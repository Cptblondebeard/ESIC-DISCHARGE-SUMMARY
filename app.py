import streamlit as st
import google.generativeai as genai
from datetime import datetime
import base64
import io
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# For PDF - using simple method that works on Replit
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor, black, white, navy

st.set_page_config(page_title="ESIC Pediatrics Discharge Summary", page_icon="🏥", layout="wide")

# Custom CSS
st.markdown("""
<style>
    .hospital-header {
        background: linear-gradient(135deg, #003366 0%, #0066CC 100%);
        padding: 20px;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 20px;
    }
    .hospital-name {
        font-size: 32px;
        font-weight: bold;
        margin-bottom: 5px;
    }
    .department-name {
        font-size: 24px;
        font-weight: 500;
        margin-bottom: 5px;
    }
    .location {
        font-size: 18px;
        opacity: 0.9;
    }
    .stButton>button {
        background: linear-gradient(135deg, #003366 0%, #0066CC 100%);
        color: white;
        font-weight: bold;
    }
    .download-section {
        background-color: #e8f4f8;
        padding: 20px;
        border-radius: 10px;
        margin-top: 20px;
        border: 2px solid #003366;
    }
</style>
""", unsafe_allow_html=True)


# Hospital Header with ESIC Logo - Properly Aligned
import os
import base64
from PIL import Image
from io import BytesIO

# Function to convert image to base64
def get_base64_of_image(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()

# Function to get image bytes for PDF
def get_image_bytes(image_path):
    with open(image_path, "rb") as img_file:
        return img_file.read()

# Check if logo file exists
logo_path = "esic_logo.png"
logo_exists = os.path.exists(logo_path)

# Main Header with Logo
if logo_exists:
    logo_base64 = get_base64_of_image(logo_path)
    
    # Create a properly aligned header with logo on left and text on right
    st.markdown(f"""
    <div style="background: linear-gradient(135deg, #003366 0%, #0066CC 100%); padding: 15px 25px; border-radius: 15px; margin-bottom: 20px; color: white; display: flex; align-items: center; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
        <div style="flex: 0 0 auto; margin-right: 25px;">
            <img src="data:image/png;base64,{logo_base64}" width="90" style="background: white; border-radius: 10px; padding: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.2);">
        </div>
        <div style="flex: 1; text-align: left;">
            <div style="font-size: 32px; font-weight: bold; margin-bottom: 5px; text-shadow: 1px 1px 2px rgba(0,0,0,0.2);">ESIC MEDICAL COLLEGE & HOSPITAL</div>
            <div style="font-size: 24px; font-weight: 500; margin-bottom: 3px; opacity: 0.95;">DEPARTMENT OF PEDIATRICS</div>
            <div style="font-size: 18px; opacity: 0.9;">KK Nagar, Chennai - 600078</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
else:
    # Fallback if logo not found
    st.markdown(f"""
    <div style="background: linear-gradient(135deg, #003366 0%, #0066CC 100%); padding: 20px; border-radius: 15px; margin-bottom: 20px; color: white; text-align: center; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
        <span style="font-size: 60px; display: block; margin-bottom: 10px;">🇮🇳</span>
        <div style="font-size: 32px; font-weight: bold; margin-bottom: 5px;">ESIC MEDICAL COLLEGE & HOSPITAL</div>
        <div style="font-size: 24px; font-weight: 500; margin-bottom: 3px;">DEPARTMENT OF PEDIATRICS</div>
        <div style="font-size: 18px;">KK Nagar, Chennai - 600078</div>
    </div>
    """, unsafe_allow_html=True)
# Initialize Gemini API
if "GEMINI_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
else:
    st.error("API Key not found. Please set GEMINI_API_KEY in Streamlit Secrets.")
# Updated to a model confirmed by your terminal test
model = genai.GenerativeModel('models/gemini-flash-latest')


# Main title with proper styling
st.markdown("""
<div style="text-align: center; padding: 15px; background: linear-gradient(135deg, #f0f6ff 0%, #e6f0ff 100%); border-radius: 12px; margin-bottom: 25px; border-left: 5px solid #003366; border-right: 5px solid #003366; box-shadow: 0 2px 4px rgba(0,0,0,0.05);">
    <h2 style="color: #003366; margin: 0; font-size: 28px;">🏥 ESIC PEDIATRICS DISCHARGE SUMMARY SYSTEM</h2>
    <p style="color: #0066CC; font-size: 16px; margin: 5px 0 0 0; font-weight: 500;">Government of India - ESIC Digital Health Initiative</p>
</div>
""", unsafe_allow_html=True)

# Create tabs
tab1, tab2, tab3 = st.tabs(["📋 Patient Details", "🔬 Clinical Data", "💊 Discharge Planning"])

with tab1:
    st.header("👤 Patient & Admission Details")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        patient_name = st.text_input("👶 Patient Name *", key="patient_name_input")
        
        # Updated Age input with years, months, and days
        st.markdown("**📅 Age ***")
        col_age1, col_age2, col_age3 = st.columns(3)
        with col_age1:
            age_years = st.number_input("Years", 0, 150, step=1, key="age_years_input")
        with col_age2:
            age_months = st.number_input("Months", 0, 11, step=1, key="age_months_input")
        with col_age3:
            age_days = st.number_input("Days", 0, 30, step=1, key="age_days_input")
        
        # Create formatted age string for display
        age_parts = []
        if age_years > 0:
            age_parts.append(f"{age_years} year{'s' if age_years > 1 else ''}")
        if age_months > 0:
            age_parts.append(f"{age_months} month{'s' if age_months > 1 else ''}")
        if age_days > 0:
            age_parts.append(f"{age_days} day{'s' if age_days > 1 else ''}")
        
        age_display = " ".join(age_parts) if age_parts else "Newborn"
        
        # Store in session state with different keys (not conflicting with widget keys)
        st.session_state['age_years_value'] = age_years
        st.session_state['age_months_value'] = age_months
        st.session_state['age_days_value'] = age_days
        st.session_state['age_display_value'] = age_display
        
        # Show age summary
        if age_years > 0 or age_months > 0 or age_days > 0:
            st.caption(f"→ {age_display}")
        
        gender = st.selectbox("⚥ Gender *", ["", "Male", "Female", "Other"], key="gender_select")
        father_name = st.text_input("👨 Father's Name", key="father_name_input")
        mother_name = st.text_input("👩 Mother's Name", key="mother_name_input")
        
    with col2:
        patient_id = st.text_input("🏥 IP Number *", key="patient_id_input")
        bed_number = st.text_input("🛏️ Bed Number", key="bed_number_input")
        unit_of_admission = st.selectbox(
            "🏛️ Unit of Admission *",
            ["", "Unit 1", "Unit 2", "Unit 3", "NICU", "PICU"],
            key="unit_select"
        )
        admission_date = st.date_input("📆 Admission Date *", key="admission_date")
        
    with col3:
        consultant_name = st.text_input("👨‍⚕️ Consultant *", key="consultant_input")
        resident_doctor = st.text_input("👨‍🔬 Resident Doctor", key="resident_input")
        discharge_date = st.date_input("📆 Discharge Date *", key="discharge_date")
        discharge_time = st.time_input("⏰ Discharge Time", value=datetime.now().time(), key="discharge_time")
    
    if admission_date and discharge_date:
        duration_of_stay = (discharge_date - admission_date).days
        st.info(f"📊 Duration of Stay: {duration_of_stay} days")
        st.markdown("---")
    # ANTHROPOMETRY SECTION
    st.subheader("📏 Anthropometry")
    a_col1, a_col2, a_col3 = st.columns(3)
    with a_col1:
        weight = st.text_input("Weight (kg)", placeholder="e.g. 10kg (50th centile)")
        height = st.text_input("Height/Length (cm)", placeholder="e.g. 75cm")
    with a_col2:
        hc = st.text_input("HC (cm)", placeholder="Head Circumference")
        muac = st.text_input("MUAC (cm)")
    with a_col3:
        wfh = st.text_input("WFH", placeholder="Weight for Height")
    
    anthro_summary = f"Weight: {weight}, Height: {height}, HC: {hc}, MUAC: {muac}, WFH: {wfh}"

with tab2:
    st.header("📋 Clinical Data")

    # NEW: Presenting Complaints (Full width at the top)
    st.subheader("🚩 Presenting Complaints *")
    presenting_complaints = st.text_area("", placeholder="Enter the symptoms that brought the patient to the hospital...", height=100, key="presenting_complaints_area")
    
    st.markdown("---")
    
    # Existing Diagnosis & Comorbidities Columns
    col_diag1, col_diag2 = st.columns(2)
    with col_diag1:
        st.subheader("📌 Admitting Diagnosis *")
        admitting_diagnosis = st.text_area("", height=100, key="admitting_diagnosis_area")
        st.subheader("📌 Comorbidities")
        comorbidities = st.text_area("", height=60, key="comorbidities_area")

    with col_diag2:
        st.subheader("✅ Discharge Diagnosis *")
        discharge_diagnosis = st.text_area("", height=100, key="discharge_diagnosis_area")
        st.subheader("📌 Complications")
        complications = st.text_area("", height=60, key="complications_area")

    st.markdown("---")
    
    st.subheader("🔬 INVESTIGATIONS")
    
    inv_col1, inv_col2 = st.columns(2)
    
    with inv_col1:
        st.markdown("### 🩸 Blood Investigations")
        blood_investigations = st.text_area(
            "Enter investigations with dates and results:",
            placeholder="Hb: 11.2 g/dL (12/03/2026)\nTLC: 15,200 cells/mm³ (12/03/2026)\nCRP: 120 mg/L (12/03/2026)",
            height=150,
            key="blood_investigations_area"
        )
        
        st.markdown("### 📊 Imaging Studies")
        imaging_investigations = st.text_area(
            "Enter imaging reports:",
            placeholder="Chest X-ray: LLL consolidation (12/03/2026)\nUSG Abdomen: Normal (13/03/2026)",
            height=150,
            key="imaging_investigations_area"
        )
        
    with inv_col2:
        st.markdown("### 🧪 Other Investigations")
        other_investigations = st.text_area(
            "Enter other tests:",
            placeholder="Urine Culture: No growth (13/03/2026)\nCSF Analysis: Normal (14/03/2026)",
            height=150,
            key="other_investigations_area"
        )
        
        st.markdown("### 📈 Vital Signs")
        vitals_trend = st.text_area(
            "Enter vital signs:",
            placeholder="BP: 110/70 mmHg\nHR: 88 bpm\nTemp: 98.6°F\nSpO2: 98%",
            height=150,
            key="vitals_trend_area"
        )
    
    st.markdown("---")
    
    st.subheader("📊 Clinical Course *")
    hospital_course = st.text_area(
        "Day-wise summary:",
        placeholder="Day 1: Admitted with fever, started on IV antibiotics\nDay 2: Improved, afebrile\nDay 3: Stable, shifted to oral\nDay 4: Discharged",
        height=150,
        key="hospital_course_area"
    )

with tab3:
    st.header("💊 Discharge Planning")
    
    col_med1, col_med2 = st.columns(2)
    
    with col_med1:
        st.subheader("💊 Discharge Medications *")
        discharge_medications = st.text_area(
            "Medications with dosage:",
            placeholder="Amoxicillin 250mg/5ml - 10ml TID x7d\nParacetamol 250mg/5ml - 10ml SOS",
            height=150,
            key="discharge_medications_area"
        )
        
        st.subheader("💉 IV Medications")
        iv_medications = st.text_area(
            "IV medications given:",
            placeholder="Ceftriaxone 500mg IV BD x3d\nIV Fluids RL",
            height=100,
            key="iv_medications_area"
        )
        
    with col_med2:
        st.subheader("📅 Follow-up Plan *")
        follow_up = st.text_area(
            "Follow-up appointments:",
            placeholder="OPD: 23/03/2026\nVaccination: MMR due\nRepeat CBC: 23/03/2026",
            height=150,
            key="follow_up_area"
        )
        
        st.subheader("⚠️ Special Instructions")
        special_instructions = st.text_area(
            "Diet, activity, precautions:",
            placeholder="Soft diet, plenty of fluids\nNo school for 1 week\nReport if fever recurs",
            height=150,
            key="special_instructions_area"
        )
    
    st.markdown("---")
    
    col_dc1, col_dc2 = st.columns(2)
    with col_dc1:
        st.subheader("✅ Discharge Condition *")
        discharge_condition = st.selectbox(
            "",
            ["", "Recovered", "Improved", "Stable", "Transferred", "LAMA", "DORB"],
            key="discharge_condition_select"
        )
        
    with col_dc2:
        st.subheader("🎯 Discharge Summary")
        discharge_advice = st.text_area(
            "Brief discharge advice:",
            placeholder="Patient discharged in stable condition. Complete antibiotic course, follow up in OPD.",
            height=100,
            key="discharge_advice_area"
        )

import os
import io
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def create_pdf_simple(summary_text, patient_name):
    buffer = io.BytesIO()
    # 1. Setup Document with proper margins
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    
    # Custom styles
    header_bold = ParagraphStyle('HeaderBold', parent=styles['Normal'], fontSize=10, leading=14, fontName='Helvetica-Bold')
    body_text = ParagraphStyle('BodyText', parent=styles['Normal'], fontSize=10, leading=14, leftIndent=12)
    
    elements = []

    # 2. Add Logo and Hospital Header correctly
    # Ensure 'esic_logo.png' is in the same folder as app.py
    logo_path = "esic_logo.png"
    if os.path.exists(logo_path):
        img = Image(logo_path, width=50, height=50)
        hospital_info = Paragraph("<b>ESIC MEDICAL COLLEGE & HOSPITAL</b><br/>Department of Pediatrics, KK Nagar, Chennai - 600078", styles['Normal'])
        header_table = Table([[img, hospital_info]], colWidths=[60, 440])
        header_table.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE')]))
        elements.append(header_table)
    else:
        # Fallback if logo is missing
        elements.append(Paragraph("<b>ESIC MEDICAL COLLEGE & HOSPITAL</b>", styles['Heading2']))
        elements.append(Paragraph("Department of Pediatrics, KK Nagar, Chennai - 600078", styles['Normal']))

    elements.append(Spacer(1, 15))
    elements.append(Paragraph(f"<u><b>DISCHARGE SUMMARY - {patient_name.upper()}</b></u>", styles['Heading3']))
    elements.append(Spacer(1, 10))

    # 3. Process the AI text line by line
    subheadings = ["NAME", "AGE", "SEX", "IP NO", "UNIT", "CONSULTANT", 
               "RESIDENT", "PRESENTING COMPLAINTS", "DIAGNOSIS", "HISTORY", 
               "ANTHROPOMETRY", "INVESTIGATIONS", "VITALS", "COURSE", 
               "TREATMENT", "ADVICE", "REVIEW", "EMERGENCY"]
    # Inside create_pdf_simple, right before the 'for line in summary_text.split' loop:
    
    # This removes any AI-generated signature lines so they don't double up
    summary_text = summary_text.replace("SIGNATURE OF THE CONSULTANT", "")
    summary_text = summary_text.replace("SIGNATURE OF THE RESIDENT", "")

    for line in summary_text.split('\n'):
        line = line.strip().replace('*', '').replace('#', '')
        if not line:
            elements.append(Spacer(1, 6))
            continue

        # --- MANUAL OVERRIDES ---
        # Update Room Number
        if "REVIEW" in line.upper():
            line = line.replace("Unit 1", "Unit 1 (Pediatric OPD Room No. 101)")
        
        # Blank out Emergency Contact
        if "EMERGENCY CONTACT" in line.upper():
            line = "<b>EMERGENCY CONTACT:</b> ________________________________"
            elements.append(Paragraph(line, styles['Normal']))
            continue # Skip the standard logic for this specific line
        # ------------------------

        # Identify if line is a Header or Body Text
        is_header = any(sub in line.upper() for sub in subheadings) and ":" in line

        if is_header:
            elements.append(Paragraph(line, header_bold))
        else:
            elements.append(Paragraph(line, body_text))


    # --- Final Signature Block ---
    elements.append(Spacer(1, 40))
    
    # We use a table to keep both signatures on one horizontal line
    sig_data = [[
        Paragraph("__________________________<br/><b>SIGNATURE OF THE CONSULTANT</b>", styles['Normal']),
        Paragraph("__________________________<br/><b>SIGNATURE OF THE RESIDENT</b>", styles['Normal'])
    ]]
    
    sig_table = Table(sig_data, colWidths=[240, 240])
    sig_table.setStyle(TableStyle([
        ('ALIGN', (0,0), (0,0), 'LEFT'),
        ('ALIGN', (1,0), (1,0), 'RIGHT'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 20),
    ]))
    
    elements.append(sig_table)

    # 5. Build and Return
    doc.build(elements)
    buffer.seek(0)
    return buffer
# WORD GENERATION with Logo
def create_word_simple(summary_text, patient_name):
    """Create Word document with ESIC logo"""
    doc = Document()
    
    # Add logo if exists
    logo_path = "esic_logo.png"
    if os.path.exists(logo_path):
        try:
            # Create a table for header with logo and text
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(1.2)
            table.columns[1].width = Inches(5)
            
            # Add logo to first cell
            cell_logo = table.cell(0, 0)
            paragraph = cell_logo.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(0.8))
            
            # Add text to second cell
            cell_text = table.cell(0, 1)
            cell_text.paragraphs[0].add_run('ESIC MEDICAL COLLEGE & HOSPITAL\n').bold = True
            cell_text.paragraphs[0].add_run('Department of Pediatrics, KK Nagar, Chennai - 600078\n')
            cell_text.paragraphs[0].add_run(f'Generated: {datetime.now().strftime("%d/%m/%Y %I:%M %p")}')
            
            doc.add_paragraph()  # Add spacing
        except:
            # Fallback if logo can't be added
            header = doc.add_heading('ESIC MEDICAL COLLEGE & HOSPITAL', 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_heading('Department of Pediatrics, KK Nagar, Chennai - 600078', 1)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%d/%m/%Y %I:%M %p")}')
    else:
        # No logo, just text
        header = doc.add_heading('ESIC MEDICAL COLLEGE & HOSPITAL', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subheader = doc.add_heading('Department of Pediatrics, KK Nagar, Chennai - 600078', 1)
        subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%d/%m/%Y %I:%M %p")}')
    
    doc.add_paragraph()
    
    # Patient name as heading
    doc.add_heading(f'Discharge Summary - {patient_name}', 2)
    doc.add_paragraph()
    
    # Add summary content with formatting
    for line in summary_text.split('\n'):
        if line.strip():
            if line.startswith('=') or line.startswith('-') or line.isupper():
                # Center alignment for section breaks
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True
            elif ':' in line and len(line.split(':')[0]) < 30:
                # Bold for labels
                p = doc.add_paragraph()
                parts = line.split(':', 1)
                runner = p.add_run(parts[0] + ':')
                runner.bold = True
                if len(parts) > 1:
                    p.add_run(parts[1])
            else:
                # Normal text
                doc.add_paragraph(line)
    
    # Add footer with logo (small)
    doc.add_paragraph()
    if os.path.exists(logo_path):
        try:
            # Add small logo in footer
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(0.3))
            paragraph.add_run('  ESIC Digital Initiative - AI Generated Discharge Summary')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            footer = doc.add_paragraph('Verified by Department of Pediatrics, ESIC Medical College, Chennai')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            footer = doc.add_paragraph('ESIC Digital Initiative - AI Generated Discharge Summary')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer = doc.add_paragraph('Verified by Department of Pediatrics, ESIC Medical College, Chennai')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        footer = doc.add_paragraph('ESIC Digital Initiative - AI Generated Discharge Summary')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer = doc.add_paragraph('Verified by Department of Pediatrics, ESIC Medical College, Chennai')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Generate button
st.markdown("---")
col_gen1, col_gen2, col_gen3 = st.columns([1, 2, 1])
with col_gen2:
    generate_btn = st.button("⚕️ GENERATE OFFICIAL DISCHARGE SUMMARY", type="primary", use_container_width=True)

# Generation logic
if generate_btn:
    # Get age values from session state
    age_years = st.session_state.get('age_years_value', 0)
    age_months = st.session_state.get('age_months_value', 0)
    age_days = st.session_state.get('age_days_value', 0)
    age_display = st.session_state.get('age_display_value', 'Newborn')
    
    # Check if age is provided (at least one field > 0)
    age_provided = (age_years > 0 or age_months > 0 or age_days > 0)
    
    if not all([patient_name, age_provided, gender, patient_id, unit_of_admission, 
                admitting_diagnosis, discharge_diagnosis, hospital_course, 
                discharge_medications, follow_up, discharge_condition]):
        st.error("⚠️ Please fill in all * marked required fields")
    else:
        with st.spinner("🧠 Gemini AI is generating official ESIC discharge summary..."):
            try:
                duration = (discharge_date - admission_date).days
                
                # We combine the system instructions and user data into one prompt for Gemini
                prompt = f"""
You are a Senior Pediatric Consultant at ESIC Medical College & Hospital. 
Create a formal discharge summary using the EXACT subheadings provided below. 
Maintain all clinical values, dates, and specific day-wise progression.
Do not add imaginary drugs, treatment or lab values. Stick to the facts 
DO NOT include any signature lines, "Signature of Consultant", or placeholder names at the end.
Stop writing immediately after the Emergency Contact section.
Donot create random phone numbers for emergency contacts 
Just Elaborate the given facts 
Do not write imaginative story
This is a medical record. Only Given facts. Just Elaborate on the facts.

--- MANDATORY STRUCTURE ---

NAME: {patient_name}
AGE: {age_display}
SEX: {gender}
IP NO: {patient_id}
UNIT: {unit_of_admission}
CONSULTANT NAME: {consultant_name}
RESIDENT NAME: {resident_doctor}
DATE OF ADMISSION: {admission_date}
DATE OF DISCHARGE: {discharge_date}
DISCHARGE DIAGNOSIS: {discharge_diagnosis}

PRESENTING COMPLAINTS: {presenting_complaints}

ADMISSION DIAGNOSIS: {admitting_diagnosis}


CLINICAL HISTORY: (Include presenting complaints, fever/respiratory details, and feeding status)
PAST HISTORY: {comorbidities if comorbidities else "None"}
ANTHROPOMETRY: {anthro_summary}
INVESTIGATIONS: (List all Blood: {blood_investigations}, Imaging: {imaging_investigations}, and Other: {other_investigations})
VITALS: {vitals_trend}
COURSE IN THE HOSPITAL: (Provide a detailed chronological narrative from admission to discharge)
TREATMENT GIVEN: (List all IV medications and significant interventions like HFNC)
DISCHARGE ADVICE: {discharge_medications} and {special_instructions}
REVIEW: {follow_up}
EMERGENCY CONTACT: 044-24891085 (Hospital) / Emergency Room

---
SIGNATURE OF THE CONSULTANT             SIGNATURE OF THE RESIDENT
"""
                
                # Gemini 1.5 Flash Call
                response = model.generate_content(
                    prompt,
                    generation_config=genai.types.GenerationConfig(
                        temperature=0.2,
                        max_output_tokens=4096
                    )
                )
                
                # Extract text directly from Gemini response
                summary = response.text
                
                # ... rest of your display and download code ...
                
                # Display summary
                st.markdown("---")
                st.success("✅ ESIC Discharge Summary Generated Successfully!")
                
                with st.container():
                    st.markdown("### 📄 OFFICIAL ESIC DISCHARGE SUMMARY")
                    st.markdown("---")
                    st.markdown(summary)
                
                # DOWNLOAD SECTION
                st.markdown('<div class="download-section">', unsafe_allow_html=True)
                st.markdown("### 📥 Download Official Documents")
                
                col_d1, col_d2, col_d3 = st.columns(3)
                
                with col_d1:
                    # Text download
                    st.download_button(
                        label="📄 Download as Text File",
                        data=summary,
                        file_name=f"ESIC_Discharge_{patient_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                        mime="text/plain",
                        use_container_width=True,
                        key="txt_download"
                    )
                
                with col_d2:
                    # PDF download
                    try:
                        pdf_buffer = create_pdf_simple(summary, patient_name)
                        st.download_button(
                            label="📕 Download as PDF",
                            data=pdf_buffer,
                            file_name=f"ESIC_Discharge_{patient_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                            mime="application/pdf",
                            use_container_width=True,
                            key="pdf_download"
                        )
                    except Exception as e:
                        st.error(f"PDF generation failed: {str(e)}")
                        st.info("Using text download as backup")
                
                with col_d3:
                    # Word download
                    try:
                        word_buffer = create_word_simple(summary, patient_name)
                        st.download_button(
                            label="📘 Download as Word",
                            data=word_buffer,
                            file_name=f"ESIC_Discharge_{patient_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key="word_download"
                        )
                    except Exception as e:
                        st.error(f"Word generation failed: {str(e)}")
                        st.info("Using text download as backup")
                
                st.markdown('</div>', unsafe_allow_html=True)
                
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")

# Sidebar
with st.sidebar:
    st.markdown("""
    <div style="text-align: center; padding: 20px; background: linear-gradient(135deg, #003366 0%, #0066CC 100%); border-radius: 10px; color: white;">        
        <h3 style="color: white;">ESIC</h3>
        <p>Employees' State Insurance Corporation</p>
        <p style="font-size: 12px;">Ministry of Labour & Employment, Govt. of India</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    st.markdown(f"⚕️ **ESIC Pediatrics System v4.2**")
    st.caption(f"Generated: {datetime.now().strftime('%d/%b/%Y %I:%M %p')}")

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; background: #f0f6ff; padding: 20px; border-radius: 10px;">
    <p style="color: #003366; font-weight: bold;">🏥 ESIC Medical College & Hospital, Department of Pediatrics, KK Nagar, Chennai - 600078</p>
    <p style="color: #666; font-size: 12px;">⚠️ This is a computer-generated discharge summary as part of ESIC's Digital Health Initiative.</p>
    <p style="color: #666; font-size: 11px;">© 2026 ESIC India - All Rights Reserved</p>
</div>

""", unsafe_allow_html=True)

